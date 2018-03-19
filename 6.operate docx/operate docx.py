#coding:utf-8

import docx

#open a document
document = docx.Document()

#add a paragraph
paragraph = document.add_paragraph('作者：苏轼', 'Subtitle')
#add a paragraph before a paragraph
prior_paragraph = paragraph.insert_paragraph_before('水调歌头','Title')

#add a long paragraph
document.add_paragraph(
'''
明月几时有，把酒问青天。 
不知天上宫阙，今夕是何年？ 
我欲乘风归去，又恐琼楼玉宇， 
高处不胜寒。 
起舞弄清影，何似在人间！ 

转朱阁，低绮户，照无眠。 
不应有恨，何事长向别时圆？ 
人有悲欢离合，月有阴晴圆缺， 
此事古难全。 
但愿人长久，千里共婵娟。
''', 'Body Text 2')

document.add_heading('Heading',)
document.add_heading('Heading', level=0)
document.add_heading('Heading', level=1)
document.add_heading('Heading', level=2)
document.save('test.docx')

